# import os
# import re
# import io
# from docx import Document  # For handling .docx files
# import uuid
# import pdfplumber
# import docx
# import pytesseract
# import cv2
# import fitz
# # from PIL import Image
# #from langchain.document_loaders import TextLoader

# # from documents_uti_oldl import (
# #     load_documents,
# #     chunk_documents,
# #     print_data_as_format_json,
# # )




# from langchain_huggingface import HuggingFaceEmbeddings
# from langchain_core.documents import Document
# from langchain_openai import ChatOpenAI



# from langchain_openai import ChatOpenAI

# from langchain_text_splitters import (
#     RecursiveCharacterTextSplitter,
#     CharacterTextSplitter,
# )

# import nltk
# from collections import Counter
# from nltk.corpus import stopwords
# from nltk.tokenize import word_tokenize
# import string
# import time

# llm = ChatOpenAI(
#     model="meta-llama/Llama-3.3-70B-Instruct-Turbo-Free",
#     openai_api_key="3cb355f55e216f29e73a674fe2d35f17438c0808eac11a726b6b364714f5871a",
#     openai_api_base="https://api.together.xyz/v1",
#     temperature=0.5
# )



# pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# import os
# import fitz
# import pytesseract
# from PIL import Image
# import pandas as pd
# import io
# import json
# from tqdm import tqdm

# def is_table_like(df):
#     if df.empty or len(df['line_num'].unique()) < 2:
#         return False
#     word_counts = df.groupby("line_num")["text"].count().values
#     return all(w >= 2 for w in word_counts[:3])

# def extract_text_and_images_with_ocr(pdf_path, output_folder):
#     os.makedirs(output_folder, exist_ok=True)
#     doc = fitz.open(pdf_path)
#     all_pages = []

#     #for page_index, page in enumerate(doc):
#     #for page_index, page in enumerate(tqdm(doc, desc="📄 Processing pages")):
#     progress = tqdm(doc, desc="📄 Processing pages")
#     for page_index, page in enumerate(progress):
#         progress.set_postfix_str(f"Page {page_index + 1}")  
#         blocks = page.get_text("dict")["blocks"]
#         page_data = []
#         used_images = set()

#         for block in blocks:
#             bbox = block["bbox"]

#             if block["type"] == 0:
#                 text = ""
#                 for line in block["lines"]:
#                     for span in line["spans"]:
#                         text += span["text"] + " "
#                 page_data.append({
#                     "type": "text",
#                     "content": text.strip(),
#                     "bbox": bbox
#                 })

#             elif block["type"] == 1:
#                 images = page.get_images(full=True)
#                 for img_index, img in enumerate(images):
#                     xref = img[0]
#                     if xref in used_images:
#                         continue
#                     used_images.add(xref)

#                     base_image = doc.extract_image(xref)
#                     image_bytes = base_image["image"]
#                     image_ext = base_image["ext"]
#                     image = Image.open(io.BytesIO(image_bytes))

#                     image_filename = f"page_{page_index + 1}_img_{img_index + 1}.{image_ext}"
#                     image_path = os.path.join(output_folder, image_filename)
#                     image.save(image_path)

#                     ocr_text = pytesseract.image_to_string(image)
#                     df = pytesseract.image_to_data(image, output_type=pytesseract.Output.DATAFRAME)
#                     df = df.dropna().query("text != ''")

#                     table_rows = []
#                     type_hint = "unknown"

#                     if is_table_like(df):
#                         type_hint = "table"
#                         for line_num in sorted(df['line_num'].unique()):
#                             line_words = df[df['line_num'] == line_num]['text'].tolist()
#                             table_rows.append(" ".join(line_words))
#                     elif "→" in ocr_text or "→" in ''.join(df["text"].values):
#                         type_hint = "flowchart"
#                     elif any(k in ocr_text.lower() for k in ["axis", "graph", "curve", "slope", "x=", "y="]):
#                         type_hint = "math_graph"
#                     elif len(ocr_text.split()) < 5:
#                         type_hint = "label_or_icon"
#                     else:
#                         type_hint = "text_image"

#                     page_data.append({
#                         "type": "image",
#                         "type_hint": type_hint,
#                         "ocr_text": ocr_text.strip(),
#                         "ocr_table": table_rows if type_hint == "table" else None,
#                         "bbox": bbox
#                     })
#                     break

#         page_data_sorted = sorted(page_data, key=lambda x: x["bbox"][1])
#         all_pages.append({
#             "page": page_index + 1,
#             "items": page_data_sorted
#         })

#     return all_pages

# def merge_text_blocks(all_pages, y_threshold=50):
#     merged = []

#     for page in all_pages:
#         text_blocks = [b for b in page["items"] if b["type"] == "text"]
#         text_blocks_sorted = sorted(text_blocks, key=lambda x: x["bbox"][1])

#         current = []
#         for block in text_blocks_sorted:
#             if not current:
#                 current.append(block)
#             else:
#                 prev = current[-1]
#                 if abs(block["bbox"][1] - prev["bbox"][3]) < y_threshold:
#                     current.append(block)
#                 else:
#                     merged.append(" ".join(b["content"] for b in current))
#                     current = [block]

#         if current:
#             merged.append(" ".join(b["content"] for b in current))

#     return merged
# def extract_text_from_pdf_v2(pdf_path):
#     output_folder = "temp_images"
#     pages_data = extract_text_and_images_with_ocr(pdf_path, output_folder)
#     merged_text_blocks = merge_text_blocks(pages_data)
#     return "\n\n".join(merged_text_blocks).strip()
# # def extract_text_from_pdf_with_ocr(pdf_path):
# #     doc = fitz.open(pdf_path)
# #     combined_text = ""

# #     for page_num in range(len(doc)):
# #         page = doc.load_page(page_num)
# #         text = page.get_text()
# #         combined_text += text + "\n"

# #         images = page.get_images(full=True)
# #         for img_index, img in enumerate(images):
# #             xref = img[0]
# #             base_image = doc.extract_image(xref)
# #             image_bytes = base_image["image"]
# #             img_ext = base_image["ext"]

# #             image = Image.open(io.BytesIO(image_bytes))
# #             ocr_text = pytesseract.image_to_string(image)
# #             combined_text += "\n[OCR IMAGE TEXT]\n" + ocr_text + "\n"

# #     return combined_text

# def extract_text_from_pdf(pdf_path):
#     text = ""
#     with pdfplumber.open(pdf_path) as pdf:
#         for page in pdf.pages:
#             text += page.extract_text() + "\n"
#     return text.strip()

# def extract_text_from_docx(docx_path):
#     doc = docx.Document(docx_path)
#     return "\n".join([para.text for para in doc.paragraphs])

# def extract_text_from_txt(txt_path):
#     with open(txt_path, "r", encoding="utf-8") as file:
#         return file.read()

# def extract_code_from_image(image_path):
#     """ Extract text from code screenshots using OCR """
#     image = cv2.imread(image_path)
#     gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)  # Convert to grayscale
#     text = pytesseract.image_to_string(gray)  # Extract text
#     return text
# def get_first_line_title(text, file_name):
#     if not text:
#         return os.path.splitext(file_name)[0]  # fallback title

#     lines = text.splitlines()
#     for line in lines:
#         clean_line = line.strip()
#         if clean_line:  # skip empty lines
#             return clean_line

#     return os.path.splitext(file_name)[0]  # fallback again if all lines are empty
# # Load all documents from a folder
# import os
# from tqdm import tqdm
# from langchain_core.documents import Document  # ✅ Correct for v1.x


# def get_first_line_title(text, file_name, max_length=120):
#     if not text:
#         return os.path.splitext(file_name)[0][:max_length]

#     lines = text.splitlines()
#     for line in lines:
#         clean_line = line.strip()
#         if clean_line and len(clean_line) <= max_length:
#             return clean_line

#     return os.path.splitext(file_name)[0][:max_length]

# def load_documents(folder_path):
#     documents = []

#     for file in tqdm(os.listdir(folder_path), desc="📂 Loading files"):
#         file_path = os.path.join(folder_path, file)

#         # 👇 Extract text first — don't append here
#         if file.endswith(".pdf"):
#             text = extract_text_from_pdf_v2(file_path)

#         elif file.endswith(".docx"):
#             text = extract_text_from_docx(file_path)

#         elif file.endswith(".txt"):
#             text = extract_text_from_txt(file_path)

#         elif file.endswith((".png", ".jpg", ".jpeg")):
#             text = extract_code_from_image(file_path)

#         else:
#             continue  # Unsupported type

#         if not text:
#             print(f"Skipping empty or unreadable file: {file}")
#             continue

#         # 👇 Build title and wrap in Document
#         title = get_first_line_title(text, file)
#         doc = Document(page_content=text, metadata={"source": file, "title": title})
#         documents.append(doc)

#     return documents

# def load_documents_v2(folder_path):
#     documents = []
#     prev_chapter = None  # Initialize prev_chapter before the loop
    
#     for file in tqdm(os.listdir(folder_path), desc="📂 Loading files"):
#         file_path = os.path.join(folder_path, file)

#         # Extract full text from each file
#         if file.endswith(".pdf"):
#             content = extract_text_from_pdf_v2(file_path)
#         elif file.endswith(".docx"):
#             content = extract_text_from_docx(file_path)
#         elif file.endswith(".txt"):
#             content = extract_text_from_txt(file_path)
#         elif file.endswith((".png", ".jpg", ".jpeg")):
#             content = extract_code_from_image(file_path)
#         else:
#             continue

#         # 📌 Detect chapter title or use fallback to previous chapter
#         chapter_title = extract_chapter_from_page(content)
        
#         # If no chapter found and we have a previous chapter, use it
#         if not chapter_title and prev_chapter:
#             chapter_title = prev_chapter
        
#         if content:
#             documents.append(Document(
#                 page_content=content,
#                 metadata={
#                     "source": file,
#                     "chapter_title": chapter_title if chapter_title else "Unknown Chapter"
#                 }
#             ))
        
#         # Update prev_chapter only if we found a new chapter
#         if chapter_title:
#             prev_chapter = chapter_title
    
#     return documents

# def extract_chapter_from_page(page_content: str) -> str:
#     """
#     Extract chapter title from a single page/document.
    
#     Looks for patterns:
#     1. "Chapter:" + separator + title (e.g., "Chapter: Microservices", "Chapter 1: Microservices")
#     2. Number + separator + title (e.g., "1. Microservices", "2 - Implementation")
    
#     Args:
#         page_content (str): Content of a single page/document
        
#     Returns:
#         str: Chapter title if found, empty string if not found
#     """
#     if not page_content or not page_content.strip():
#         return ""
    
#     lines = page_content.split('\n')
    
#     for line in lines:
#         line = line.strip()
#         if not line:
#             continue
        
#         # Pattern 1: "Chapter" + optional number + separator + title
#         # Matches: "Chapter: Microservices", "Chapter 1: Microservices", "Chapter - Implementation"
#         chapter_pattern = r'^\s*chapter\s*(?:\d+)?\s*[\:\-\.\s]+(.+?)$'
#         match = re.match(chapter_pattern, line, re.IGNORECASE)
#         if match:
#             title = match.group(1).strip()
#             if title:
#                 return title
        
#         # Pattern 2: Whole number + separator + title (NOT decimal numbers like 1.1, 1.22)
#         # Matches: "1. Microservices", "2 - Implementation", "3: Use Cases", "1 Microservices"
#         # But NOT: "1.1 What are...", "1.22 Details..."
#         number_pattern = r'^\s*(\d+)\s*[\:\-\s]+(.+?)$|^\s*(\d+)\.(?!\d)\s*(.+?)$'
#         match = re.match(number_pattern, line, re.IGNORECASE)
#         if match:
#             # Handle both pattern groups (with different separators)
#             title = (match.group(2) or match.group(4)).strip() if match.group(2) or match.group(4) else ""
#             if title:
#                 return title
    
#     return ""



# nltk.download("punkt_tab")  # Changed from "punkt" to "punkt_tab"
# nltk.download("stopwords")


# # fkind first chapter /couser/lesson title to the chunke
# import re


# def find_first_chapter_title_v2(text, mode="chapter", max_line_length=120):
#     lines = text.splitlines()

#     for line in lines:
#         line = line.strip()
#         if not line:
#             continue

#         # 🔍 Step 1: Try known chapter patterns
#         if mode in ("chapter", "flexible"):
#             patterns = [
#                 r"(?i)\bchapter\s*\d+[:\-\.\s]*(.+)",
#                 r"(?i)\b\d+\.\s+(.+)",
#                 r"(?i)\bcourse[:\-\.\s]*(.+)",
#                 r"(?i)\blesson\s*\d+[:\-\.\s]*(.+)",
#                 r"(?i)^(.+):\s*$"  # Line ends in colon with a title
#             ]

#             for pattern in patterns:
#                 match = re.search(pattern, line)
#                 if match:
#                     return match.group(1).strip()

#         # 🧪 Step 2: Use first non-empty line (with or without `:`)
#         if mode in ("first_line", "flexible"):
#             if len(line) <= max_line_length:
#                 return line.rstrip(":").strip()

#     return None







# def extract_keywords(text, top_n=5):
#     stop_words = set(stopwords.words("english"))
#     words = word_tokenize(text.lower())
#     words = [word for word in words if word.isalnum() and word not in stop_words]
    
#     most_common = Counter(words).most_common(top_n)
#     keywords = [word for word, _ in most_common]
#     return keywords



# # here i used manual semantic chunking by: page or logic  logic separator,summary from prev chunk,keyword
# import uuid
# import time

# # def chunk_documents(docs, chunk_size=500, chunk_overlap=50, page=False, summary_chunk=False, keywords_chunk=False):
# #     chunked_docs = []
# #     ## if page is as true the give prioity for page then logical separators
# #      ## if page is as false the give the logical separators

# #     separators = ["page", "\n\n", "\n", ".", " ", ""] if page else ["\n\n", "\n", ".", " ", ""]

# #     text_splitter = RecursiveCharacterTextSplitter(
# #         separators=separators,
# #         chunk_size=chunk_size,
# #         chunk_overlap=chunk_overlap
# #     )

# #     if isinstance(docs, str):
# #         docs = [Document(page_content=docs, metadata={})]  # Wrap raw text in a Document

# #     if isinstance(docs, list) and all(isinstance(doc, Document) for doc in docs):
# #         chunks = text_splitter.split_documents(docs)
# #     else:
# #         docs = [Document(page_content=str(doc), metadata={}) for doc in docs]
# #         chunks = text_splitter.split_documents(docs)

# #     prev_summary = None
# #     prev_keywords = None

# #     for idx, chunk in enumerate(chunks):
# #         if not chunk.page_content.strip():
# #             continue

# #         summary = (
# #             llm.invoke(f"Summarize this text in 1 sentence:\n\n{chunk.page_content}").content
# #             if summary_chunk else ""
# #         )

# #         keywords = extract_keywords(chunk.page_content) if keywords_chunk else ""

# #         chapter_title = find_first_chapter_title_v2(chunk.page_content)
        
# #         #  If found, update the current chapter
# #         if chapter_title:
# #             current_chapter = chapter_title

# #         # Get original source and title from the metadata
# #         source = chunk.metadata.get("source", "unknown_source")
# #         title = chunk.metadata.get("title", "unknown_title")
# #         chapter_title=title
# #         chunked_docs.append({
# #             "id": str(uuid.uuid4()),
# #             "chapter": current_chapter,
# #             "source": source,
# #             "title": title,
# #             "text": chunk.page_content,
# #             "metadata": {
# #                 "summary": summary,
# #                 "prev_summary": prev_summary,
# #                 "chunk_index": idx,
# #                 "chunk_length": len(chunk.page_content),
# #                 "prev_keywords": prev_keywords,
# #                 "keywords": keywords
# #             }
# #         })

# #         prev_summary = summary
# #         prev_keywords = keywords

# #     return chunked_docs

# def chunk_documents(docs, chunk_size=500, chunk_overlap=50, page=False, summary_chunk=False, keywords_chunk=False, cv_mode=False):
#     chunked_docs = []
    
#     # CV mode: prioritize line breaks to keep logical sections intact
#     if cv_mode:
#         separators = ["\n\n", "\n", ".", " ", ""]
#         # Reduce overlap for CVs since sections are typically independent
#         chunk_overlap = min(chunk_overlap, 20)
#     elif page:
#         separators = ["page", "\n\n", "\n", ".", " ", ""]
#     else:
#         separators = ["\n\n", "\n", ".", " ", ""]

#     text_splitter = RecursiveCharacterTextSplitter(
#         separators=separators,
#         chunk_size=chunk_size,
#         chunk_overlap=chunk_overlap
#     )

#     if isinstance(docs, str):
#         docs = [Document(page_content=docs, metadata={})]

#     if isinstance(docs, list) and all(isinstance(doc, Document) for doc in docs):
#         chunks = text_splitter.split_documents(docs)
#     else:
#         docs = [Document(page_content=str(doc), metadata={}) for doc in docs]
#         chunks = text_splitter.split_documents(docs)

#     prev_summary = None
#     prev_keywords = None
#     current_chapter = "Unknown Chapter"

#     for idx, chunk in enumerate(chunks):
#         if not chunk.page_content.strip():
#             continue

#         summary = (
#             llm.invoke(f"Summarize this text in 1 sentence:\n\n{chunk.page_content}").content
#             if summary_chunk else ""
#         )

#         keywords = extract_keywords(chunk.page_content) if keywords_chunk else ""

#         # For CVs, you might want to detect section headers differently
#         if cv_mode:
#             chapter_title = detect_cv_section(chunk.page_content)
#         else:
#             chapter_title = find_first_chapter_title_v2(chunk.page_content)
        
#         if chapter_title:
#             current_chapter = chapter_title

#         source = chunk.metadata.get("source", "unknown_source")
#         title = chunk.metadata.get("title", "unknown_title")
        
#         chunked_docs.append({
#             "id": str(uuid.uuid4()),
#             "chapter": current_chapter,
#             "source": source,
#             "title": title,
#             "text": chunk.page_content,
#             "metadata": {
#                 "summary": summary,
#                 "prev_summary": prev_summary,
#                 "chunk_index": idx,
#                 "chunk_length": len(chunk.page_content),
#                 "prev_keywords": prev_keywords,
#                 "keywords": keywords
#             }
#         })

#         prev_summary = summary
#         prev_keywords = keywords

#     return chunked_docs


# def detect_cv_section(text):
#     """Detect common CV section headers"""
#     cv_sections = [
#         "experience", "work experience", "education", "skills", 
#         "projects", "certifications", "summary", "objective",
#         "languages", "awards", "publications", "contact"
#     ]
    
#     first_line = text.split('\n')[0].lower().strip()
    
#     for section in cv_sections:
#         if section in first_line:
#             return first_line.title()
    
#     return None        

# import json
# import pprint

# def print_data_as_format_json(docs):
#     pp = pprint.PrettyPrinter(
#         indent=2, 
#         width=120,        # Wider output
#         depth=None,       # No depth limit
#         compact=False     # Don't compact
#     )
#     pp.pprint(docs)

"""
document_processor.py
─────────────────────
Improved version with:
- TOC page detection and skipping
- Smart chunking per document type (PDF / notes / profile)
- Clean chapter name extraction (no TOC lines)
- Minimum chunk length filter
- Recursive subfolder loading
"""

import os
import re
import io
import uuid
import json
import pprint
import string
import time

import cv2
import fitz
import docx
import pdfplumber
import pytesseract
import nltk
from PIL import Image
from tqdm import tqdm
from collections import Counter

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_openai import ChatOpenAI
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize

# ── LLM (for optional summary) ───────────────────────────────
llm = ChatOpenAI(
    model="meta-llama/Llama-3.3-70B-Instruct-Turbo-Free",
    openai_api_key="3cb355f55e216f29e73a674fe2d35f17438c0808eac11a726b6b364714f5871a",
    openai_api_base="https://api.together.xyz/v1",
    temperature=0.5
)

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

nltk.download("punkt_tab", quiet=True)
nltk.download("stopwords", quiet=True)

# ── Constants ─────────────────────────────────────────────────
MIN_CHUNK_LENGTH = 100       # skip chunks shorter than this
MAX_CHAPTER_LENGTH = 80      # cap chapter name length


# ═══════════════════════════════════════════════════════════════
# TOC DETECTION
# ═══════════════════════════════════════════════════════════════

def is_toc_line(line: str) -> bool:
    """Detect a single TOC line: 'Chapter name.......8'"""
    return bool(re.search(r'\.{3,}\s*\d{1,3}\s*$', line.strip()))


def is_toc_block(text: str, threshold: int = 3) -> bool:
    """
    Return True if this text block looks like a Table of Contents.
    A block with 3+ TOC-style lines is considered a TOC page.
    """
    lines = text.splitlines()
    toc_line_count = sum(1 for l in lines if is_toc_line(l))
    return toc_line_count >= threshold


def is_watermark_line(line: str) -> bool:
    """Detect page headers, footers, and watermark lines."""
    patterns = [
        r"^TheAIEngineers\s*\d*$",
        r"^Page\s+\d+\s+of\s+\d+$",
        r"^\d+\s*$",
        r"^Check the complete article",
        r"^www\.",
        r"^http",
        r"^©",
        r"^All rights reserved",
    ]
    line = line.strip()
    return any(re.match(p, line, re.IGNORECASE) for p in patterns)


def is_merged_toc_block(text: str) -> bool:
    """
    Detect TOC merged into one block without dot leaders.
    Pattern: multiple 'title ... number' sequences.
    """
    matches = re.findall(r'[A-Za-z].+?\s+\d{1,3}\s{2,}', text)
    return len(matches) >= 3


def filter_toc_from_text(text: str) -> str:
    """
    Remove TOC lines and watermark lines from extracted text.
    Also removes merged TOC blocks.
    """
    lines = text.splitlines()
    filtered = []
    for line in lines:
        if is_toc_line(line):
            continue
        if is_watermark_line(line):
            continue
        filtered.append(line)
    return "\n".join(filtered).strip()


# ═══════════════════════════════════════════════════════════════
# CHAPTER / TITLE DETECTION
# ═══════════════════════════════════════════════════════════════

def clean_chapter_name(chapter: str) -> str:
    """Remove TOC dot patterns and page numbers from chapter names."""
    clean = re.sub(r'\.{2,}\s*\d+', '', chapter)   # remove "....8"
    clean = re.sub(r'\s{2,}', ' ', clean).strip()
    return clean[:MAX_CHAPTER_LENGTH] if clean else ""


def find_first_chapter_title_v2(text: str, mode: str = "flexible") -> str:
    """
    Extract chapter/section title from a chunk.
    Skips TOC lines and overly long lines.
    """
    lines = text.splitlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Skip TOC lines and watermarks
        if is_toc_line(line):
            continue
        if is_watermark_line(line):
            continue

        # Skip very long lines (likely TOC entries or paragraph text)
        if len(line) > MAX_CHAPTER_LENGTH:
            continue

        if mode in ("chapter", "flexible"):
            patterns = [
                r"(?i)^chapter\s*\d+[:\-\.\s]+(.+)",
                r"(?i)^\d+\.\d+[\d\.]*\s+(.+)",   # "2.1 What is it about?" "5.1.1 Details"
                r"(?i)^\d+\s+([A-Z][a-z].+)",          # "1 Prerequisites" "2 Introduction"
                r"(?i)^\d+\.\s+([A-Z].+)",            # "1. Introduction"
                r"(?i)^course[:\-\.\s]+(.+)",
                r"(?i)^lesson\s*\d+[:\-\.\s]+(.+)",
                r"(?i)^(introduction|conclusion|summary|overview|background|prerequisites)$",
            ]
            for pattern in patterns:
                match = re.match(pattern, line)
                if match:
                    groups = match.groups()
                    title = groups[-1].strip() if groups[-1] else line.strip()
                    return clean_chapter_name(title)

        if mode in ("first_line", "flexible"):
            # Use first short non-empty non-TOC line as title
            candidate = line.rstrip(":").strip()
            if candidate:
                return clean_chapter_name(candidate)

    return ""


def detect_cv_section(text: str) -> str:
    """Detect common CV/profile section headers."""
    cv_sections = [
        "experience", "work experience", "education", "skills",
        "projects", "certifications", "summary", "objective",
        "languages", "awards", "publications", "contact", "profile"
    ]
    first_line = text.split('\n')[0].lower().strip()
    for section in cv_sections:
        if section in first_line:
            return first_line.title()
    return ""


def get_first_line_title(text: str, file_name: str, max_length: int = 120) -> str:
    """Extract document title from first non-empty non-watermark line."""
    if not text:
        return os.path.splitext(file_name)[0][:max_length]
    for line in text.splitlines():
        clean = line.strip()
        if (clean
                and len(clean) <= max_length
                and not is_toc_line(clean)
                and not is_watermark_line(clean)):
            return clean
    return os.path.splitext(file_name)[0][:max_length]


# ═══════════════════════════════════════════════════════════════
# DOCUMENT TYPE DETECTION
# ═══════════════════════════════════════════════════════════════

def detect_doc_type(file_name: str, text: str = "") -> str:
    """
    Classify document into type for smart chunking.

    Returns: 'profile' | 'pdf_course' | 'notes'
    """
    name_lower = file_name.lower()

    if "profile" in name_lower or "cv" in name_lower or "resume" in name_lower:
        return "profile"

    if file_name.endswith(".pdf"):
        return "pdf_course"

    # .txt / .md / .docx = personal notes
    return "notes"


def get_chunk_settings(doc_type: str) -> dict:
    """
    Return chunking settings per document type.

    profile    → large chunks, keep context together
    pdf_course → medium chunks, filter TOC
    notes      → smaller chunks, clean structure
    """
    settings = {
        "profile": {
            "chunk_size": 1500,
            "chunk_overlap": 200,
            "cv_mode": True,
        },
        "pdf_course": {
            "chunk_size": 1000,
            "chunk_overlap": 150,
            "cv_mode": False,
        },
        "notes": {
            "chunk_size": 800,
            "chunk_overlap": 100,
            "cv_mode": False,
        },
    }
    return settings.get(doc_type, settings["notes"])


# ═══════════════════════════════════════════════════════════════
# TEXT EXTRACTION
# ═══════════════════════════════════════════════════════════════

def is_table_like(df) -> bool:
    if df.empty or len(df['line_num'].unique()) < 2:
        return False
    word_counts = df.groupby("line_num")["text"].count().values
    return all(w >= 2 for w in word_counts[:3])


def extract_text_and_images_with_ocr(pdf_path: str, output_folder: str) -> list:
    os.makedirs(output_folder, exist_ok=True)
    doc = fitz.open(pdf_path)
    all_pages = []

    progress = tqdm(doc, desc="Processing pages")
    for page_index, page in enumerate(progress):
        progress.set_postfix_str(f"Page {page_index + 1}")
        blocks = page.get_text("dict")["blocks"]
        page_data = []
        used_images = set()

        for block in blocks:
            bbox = block["bbox"]

            if block["type"] == 0:
                text = ""
                for line in block["lines"]:
                    for span in line["spans"]:
                        text += span["text"] + " "
                page_data.append({
                    "type": "text",
                    "content": text.strip(),
                    "bbox": bbox
                })

            elif block["type"] == 1:
                images = page.get_images(full=True)
                for img_index, img in enumerate(images):
                    xref = img[0]
                    if xref in used_images:
                        continue
                    used_images.add(xref)

                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    image = Image.open(io.BytesIO(image_bytes))

                    image_filename = f"page_{page_index + 1}_img_{img_index + 1}.{image_ext}"
                    image_path = os.path.join(output_folder, image_filename)
                    image.save(image_path)

                    ocr_text = pytesseract.image_to_string(image)
                    df = pytesseract.image_to_data(image, output_type=pytesseract.Output.DATAFRAME)
                    df = df.dropna().query("text != ''")

                    table_rows = []
                    type_hint = "unknown"

                    if is_table_like(df):
                        type_hint = "table"
                        for line_num in sorted(df['line_num'].unique()):
                            line_words = df[df['line_num'] == line_num]['text'].tolist()
                            table_rows.append(" ".join(line_words))
                    elif "→" in ocr_text:
                        type_hint = "flowchart"
                    elif any(k in ocr_text.lower() for k in ["axis", "graph", "curve", "x=", "y="]):
                        type_hint = "math_graph"
                    elif len(ocr_text.split()) < 5:
                        type_hint = "label_or_icon"
                    else:
                        type_hint = "text_image"

                    page_data.append({
                        "type": "image",
                        "type_hint": type_hint,
                        "ocr_text": ocr_text.strip(),
                        "ocr_table": table_rows if type_hint == "table" else None,
                        "bbox": bbox
                    })
                    break

        page_data_sorted = sorted(page_data, key=lambda x: x["bbox"][1])
        all_pages.append({"page": page_index + 1, "items": page_data_sorted})

    return all_pages


def merge_text_blocks(all_pages: list, y_threshold: int = 50) -> list:
    merged = []
    for page in all_pages:
        text_blocks = sorted(
            [b for b in page["items"] if b["type"] == "text"],
            key=lambda x: x["bbox"][1]
        )
        current = []
        for block in text_blocks:
            if not current:
                current.append(block)
            else:
                prev = current[-1]
                if abs(block["bbox"][1] - prev["bbox"][3]) < y_threshold:
                    current.append(block)
                else:
                    merged.append(" ".join(b["content"] for b in current))
                    current = [block]
        if current:
            merged.append(" ".join(b["content"] for b in current))
    return merged


def extract_text_from_pdf_v2(pdf_path: str) -> str:
    pages_data = extract_text_and_images_with_ocr(pdf_path, "temp_images")
    merged = merge_text_blocks(pages_data)
    full_text = "\n\n".join(merged).strip()
    # Remove TOC lines from full text
    return filter_toc_from_text(full_text)


def extract_text_from_pdf(pdf_path: str) -> str:
    text = ""
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"
    return filter_toc_from_text(text.strip())


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text from .docx preserving heading structure.
    Heading paragraphs get extra newlines so chapter detection works.
    TOC lines are filtered out automatically.
    """
    document = docx.Document(docx_path)
    lines = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Skip TOC lines
        if is_toc_line(text):
            continue
        # Preserve heading levels with surrounding newlines
        if para.style.name.startswith("Heading"):
            lines.append(f"\n{text}\n")
        else:
            lines.append(text)
    return "\n".join(lines).strip()


def extract_text_from_txt(txt_path: str) -> str:
    with open(txt_path, "r", encoding="utf-8") as f:
        return f.read()


def extract_code_from_image(image_path: str) -> str:
    image = cv2.imread(image_path)
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    return pytesseract.image_to_string(gray)


# ═══════════════════════════════════════════════════════════════
# DOCUMENT LOADING
# ═══════════════════════════════════════════════════════════════

def load_documents(folder_path: str) -> list:
    """
    Load all supported documents from folder and subfolders.
    Filters TOC content from PDFs automatically.
    """
    documents = []
    supported = (".pdf", ".docx", ".doc", ".txt", ".md", ".png", ".jpg", ".jpeg")

    # Collect all files recursively
    all_files = []
    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith(supported):
                all_files.append((root, file))

    for root, file in tqdm(all_files, desc="Loading files"):
        file_path = os.path.join(root, file)
        text = None

        try:
            if file.lower().endswith(".pdf"):
                text = extract_text_from_pdf_v2(file_path)
            elif file.lower().endswith((".docx", ".doc")):
                text = extract_text_from_docx(file_path)
            elif file.lower().endswith((".txt", ".md")):
                text = extract_text_from_txt(file_path)
            elif file.lower().endswith((".png", ".jpg", ".jpeg")):
                text = extract_code_from_image(file_path)
        except Exception as e:
            print(f"Error loading {file}: {e}")
            continue

        if not text or not text.strip():
            print(f"Skipping empty: {file}")
            continue

        relative_path = os.path.relpath(file_path, folder_path)
        title = get_first_line_title(text, file)
        doc_type = detect_doc_type(file, text)

        doc = Document(
            page_content=text,
            metadata={
                "source": relative_path,
                "title": title,
                "doc_type": doc_type,
                "file_name": file,
            }
        )
        documents.append(doc)

    print(f"\nLoaded {len(documents)} documents")
    return documents


# ═══════════════════════════════════════════════════════════════
# KEYWORDS
# ═══════════════════════════════════════════════════════════════

def extract_keywords(text: str, top_n: int = 5) -> list:
    stop_words = set(stopwords.words("english"))
    words = word_tokenize(text.lower())
    words = [w for w in words if w.isalnum() and w not in stop_words]
    return [word for word, _ in Counter(words).most_common(top_n)]


# ═══════════════════════════════════════════════════════════════
# SMART CHUNKING
# ═══════════════════════════════════════════════════════════════

def chunk_documents(
    docs,
    chunk_size: int = None,
    chunk_overlap: int = None,
    page: bool = False,
    summary_chunk: bool = False,
    keywords_chunk: bool = False,
    cv_mode: bool = False,
    smart_mode: bool = True,      # use per-doc-type settings
) -> list:
    """
    Chunk documents with smart per-type settings.

    smart_mode=True  → uses doc_type metadata to pick chunk size
    smart_mode=False → uses provided chunk_size/chunk_overlap
    """
    if isinstance(docs, str):
        docs = [Document(page_content=docs, metadata={})]
    elif not isinstance(docs[0], Document):
        docs = [Document(page_content=str(d), metadata={}) for d in docs]

    chunked_docs = []

    for doc in docs:
        doc_type = doc.metadata.get("doc_type", "notes")

        # Determine settings
        if smart_mode:
            settings = get_chunk_settings(doc_type)
            c_size    = settings["chunk_size"]
            c_overlap = settings["chunk_overlap"]
            is_cv     = settings["cv_mode"]
        else:
            c_size    = chunk_size or 1000
            c_overlap = chunk_overlap or 150
            is_cv     = cv_mode

        separators = ["\n\n", "\n", ".", " ", ""]
        if page:
            separators = ["page"] + separators

        splitter = RecursiveCharacterTextSplitter(
            separators=separators,
            chunk_size=c_size,
            chunk_overlap=c_overlap,
        )

        chunks = splitter.split_documents([doc])

        prev_summary  = None
        prev_keywords = None
        current_chapter = doc.metadata.get("title", "General")

        for idx, chunk in enumerate(chunks):
            text = chunk.page_content.strip()

            # Skip empty or too-short chunks
            if not text or len(text) < MIN_CHUNK_LENGTH:
                continue

            # Skip TOC blocks that slipped through
            if is_toc_block(text):
                continue

            # Skip merged TOC blocks (no dot leaders)
            if is_merged_toc_block(text):
                continue

            # Chapter detection
            if is_cv:
                detected = detect_cv_section(text)
            else:
                detected = find_first_chapter_title_v2(text)

            if detected:
                current_chapter = detected

            summary = (
                llm.invoke(f"Summarize this text in 1 sentence:\n\n{text}").content
                if summary_chunk else ""
            )
            keywords = extract_keywords(text) if keywords_chunk else []

            source = chunk.metadata.get("source", "unknown")
            title  = chunk.metadata.get("title", "unknown")

            chunked_docs.append({
                "id": str(uuid.uuid4()),
                "chapter": current_chapter,
                "source": source,
                "title": title,
                "doc_type": doc_type,
                "text": text,
                "metadata": {
                    "summary": summary,
                    "prev_summary": prev_summary,
                    "chunk_index": idx,
                    "chunk_length": len(text),
                    "prev_keywords": prev_keywords,
                    "keywords": keywords,
                }
            })

            prev_summary  = summary
            prev_keywords = keywords

    print(f"Created {len(chunked_docs)} chunks from {len(docs)} documents")
    return chunked_docs


# ═══════════════════════════════════════════════════════════════
# DEBUG UTILS
# ═══════════════════════════════════════════════════════════════

def print_data_as_format_json(docs):
    pp = pprint.PrettyPrinter(indent=2, width=120, depth=None, compact=False)
    pp.pprint(docs)


def print_chunk_stats(chunked_docs: list):
    """Print summary statistics about the chunked documents."""
    if not chunked_docs:
        print("No chunks to analyze")
        return

    lengths = [c["metadata"]["chunk_length"] for c in chunked_docs]
    doc_types = {}
    for c in chunked_docs:
        dt = c.get("doc_type", "unknown")
        doc_types[dt] = doc_types.get(dt, 0) + 1

    print("\n" + "=" * 50)
    print("CHUNK STATISTICS")
    print("=" * 50)
    print(f"Total chunks:     {len(chunked_docs)}")
    print(f"Avg chunk length: {sum(lengths) // len(lengths)} chars")
    print(f"Min chunk length: {min(lengths)} chars")
    print(f"Max chunk length: {max(lengths)} chars")
    print(f"\nChunks by doc type:")
    for dt, count in doc_types.items():
        print(f"  {dt}: {count} chunks")
    print("=" * 50)
